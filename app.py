# app.py
import streamlit as st
import requests
from rembg import remove
from PIL import Image
import os
from geopy.geocoders import Nominatim
from streamlit_js_eval import get_geolocation
import folium
from streamlit_folium import st_folium
import pdfplumber
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from datetime import datetime
from io import BytesIO
import zipfile
import pyperclip
from st_copy_to_clipboard import st_copy_to_clipboard
from streamlit_image_comparison import image_comparison
from constancia import (combinar_pdfs, configure_selenium_driver, download_rnp_certificate, download_sunat_ruc_pdf, download_rnssc_pdf, descargar_constancias, setup_logging, add_success_method)
import logging
logging.basicConfig(
level=logging.INFO, 
format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('constancia')
def verificar_configuracion_selenium():
    """
    Diagnóstico completo de la configuración de Selenium
    """
    try:
        # Verificar dependencias
        import selenium
        import webdriver_manager
        
        logger.info(f"Versión de Selenium: {selenium.__version__}")
        logger.info(f"Versión de WebDriver Manager: {webdriver_manager.__version__}")
        
        # Verificar variables de entorno
        logger.info("Variables de entorno:")
        logger.info(f"PATH: {os.environ.get('PATH')}")
        logger.info(f"HOME: {os.environ.get('HOME')}")
        
        # Verificar directorios de trabajo
        logger.info("Directorios:")
        logger.info(f"Directorio actual: {os.getcwd()}")
        logger.info(f"Contenido del directorio: {os.listdir('.')}")
        
    except Exception as e:
        logger.error(f"Error en verificación de configuración: {e}")

def descargar_constancias_debug(ruc, dni):
    """
    Versión de depuración de la descarga de constancias
    """
    try:
        # Crear directorio de trabajo
        output_dir = '/mount/src/cotizacionprueba'
        os.makedirs(output_dir, exist_ok=True)
        
        # Logging extensivo
        logger.debug(f"Intentando descargar constancias para RUC: {ruc}, DNI: {dni}")
        logger.debug(f"Directorio de descargas: {os.path.abspath(output_dir)}")
        
        # Verificar permisos
        logger.debug(f"Permisos del directorio: {oct(os.stat(output_dir).st_mode)[-3:]}")
        
        # Intentar descargas individuales con más logging
        try:
            # Descarga RNP
            logger.debug("Iniciando descarga RNP")
            rnp_file = download_rnp_certificate(ruc, output_dir)
            logger.debug(f"Resultado descarga RNP: {rnp_file}")
        except Exception as e:
            logger.error(f"Error en descarga RNP: {e}", exc_info=True)
        
        try:
            # Descarga RUC
            logger.debug("Iniciando descarga RUC")
            ruc_file = download_sunat_ruc_pdf(ruc, output_dir)
            logger.debug(f"Resultado descarga RUC: {ruc_file}")
        except Exception as e:
            logger.error(f"Error en descarga RUC: {e}", exc_info=True)
        
        try:
            # Descarga RNSSC
            logger.debug("Iniciando descarga RNSSC")
            rnssc_file = download_rnssc_pdf(dni, output_dir)
            logger.debug(f"Resultado descarga RNSSC: {rnssc_file}")
        except Exception as e:
            logger.error(f"Error en descarga RNSSC: {e}", exc_info=True)
        
        # Verificar archivos descargados
        archivos = os.listdir(output_dir)
        logger.debug(f"Archivos descargados: {archivos}")
        
        return archivos
    
    except Exception as e:
        logger.error(f"Error general en descarga de constancias: {e}", exc_info=True)
        return None

# Determinar la ruta base de la aplicación
base_dir = os.path.dirname(os.path.abspath(__file__))

def obtener_datos_sunat(dni):
    apisnet_key = st.secrets["APISNET"]["key"]
    url = f"https://api.apis.net.pe/v2/sunat/dni?numero={dni}&token={apisnet_key}"
    try:
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            nombres = f"{data.get('nombres', '')} {data.get('apellidoPaterno', '')} {data.get('apellidoMaterno', '')}".strip()
            ruc = data.get("ruc", "")
            return nombres, ruc
        else:
            st.error("Error al obtener datos de SUNAT. Verifica el DNI ingresado.")
            return None, None
    except Exception as e:
        st.error(f"Error al conectar con la API de SUNAT: {e}")
        return None, None

def obtener_direccion_desde_coordenadas(lat, lon):
    geolocator = Nominatim(user_agent="my_streamlit_app")
    try:
        location = geolocator.reverse((lat, lon))
        return location.address
    except Exception as e:
        st.error(f"Error al obtener la dirección: {e}")
        return None

def crear_mapa(lat=None, lon=None, zoom=13):
    # Usar coordenadas proporcionadas o predeterminadas de Lima, Perú
    if lat is None or lon is None:
        lat, lon = -12.0464, -77.0428  # Coordenadas de Lima

    # Crear el mapa centrado en la ubicación y usando el zoom proporcionado
    m = folium.Map(location=[lat, lon], zoom_start=zoom)

    # Agregar marcador si hay coordenadas específicas
    folium.Marker(
        [lat, lon],
        popup="Ubicación actual",
        icon=folium.Icon(color='red', icon='info-sign'),
        draggable=False
    ).add_to(m)

    return m

def extraer_nombre_servicio(pdf_file):
    texto_completo = ""
    with pdfplumber.open(pdf_file) as pdf:
        for pagina in pdf.pages:
            texto_completo += pagina.extract_text()

    texto_unido = ' '.join(texto_completo.split())

    patron = r'2\.\s*OBJETO\s*DE\s*LA\s*CONTRATACION\s*(.*?)\s*3\.\s*FINALIDAD\s*PUBLICA'

    match = re.search(patron, texto_unido, re.DOTALL | re.IGNORECASE)

    if match:
        servicio = ' '.join(match.group(1).split())
        return servicio
    return "Servicio no encontrado"

def extraer_forma_pago(pdf_file):
    texto_completo = ""
    with pdfplumber.open(pdf_file) as pdf:
        for pagina in pdf.pages:
            texto_completo += pagina.extract_text()

    texto_unido = ' '.join(texto_completo.split())

    patron = r'El pago se realizará en\s*(.*?)\s*luego de la emisión de la conformidad del servicio,'

    match = re.search(patron, texto_unido, re.DOTALL | re.IGNORECASE)

    if match:
        forma_pago = ' '.join(match.group(1).split()).upper()
        return forma_pago
    return "FORMA DE PAGO NO ENCONTRADA"

def extraer_dias(pdf_file):
    texto_completo = ""
    with pdfplumber.open(pdf_file) as pdf:
        for pagina in pdf.pages:
            texto_completo += pagina.extract_text()

    texto_unido = ' '.join(texto_completo.split())

    patron = r'El plazo de ejecución del servicio es de hasta\s*(\d+)\s*días calendario'

    match = re.search(patron, texto_unido, re.DOTALL | re.IGNORECASE)

    if match:
        dias = match.group(1)
        return dias
    return "DÍAS NO ENCONTRADOS"

def obtener_valor_sugerido(dias):
    """
    Determina el valor sugerido basado en los días de ejecución
    """
    try:
        dias = int(dias)
        if dias <= 30:
            return 2000.0
        elif dias <= 60:
            return 4000.0
        elif dias <= 90:
            return 6000.0
        elif dias <= 120:
            return 8000.0
        else:
            return 8000.0  # Valor máximo por defecto
    except (ValueError, TypeError):
        return 2000.0  # Valor por defecto si hay error en la conversión

def procesar_firma(firma_file, remover_fondo=False):
    """
    Procesa la imagen de la firma, opcionalmente removiendo el fondo.
    
    Args:
        firma_file: Archivo de imagen subido
        remover_fondo: Boolean indicando si se debe remover el fondo
    
    Returns:
        BytesIO: Imagen procesada en formato BytesIO
    """
    # Abrir la imagen
    image = Image.open(firma_file)
    
    if remover_fondo:
        with st.spinner('Removiendo fondo de la firma...'):
            # Remover fondo
            imagen_procesada = remove(image)
            # Convertir a modo RGBA si no lo está ya
            if imagen_procesada.mode != 'RGBA':
                imagen_procesada = imagen_procesada.convert('RGBA')
    else:
        imagen_procesada = image
        
    # Convertir a BytesIO
    img_byte_arr = BytesIO()
    imagen_procesada.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)
    
    return img_byte_arr

def mostrar_seccion_firma():
    """
    Muestra la sección de carga y procesamiento de firma
    
    Returns:
        BytesIO: Imagen de firma procesada
        bool: Indicador de si se cargó una firma
    """
    st.header("Sube tu firma")
    
    # Checkbox para remover fondo
    remover_fondo = st.checkbox("Remover fondo de la firma", value=False)
    
    # Upload de firma
    firma_file = st.file_uploader(
        "Selecciona tu imagen de firma", 
        type=["png", "jpg", "jpeg"],
        help="Sube una imagen de tu firma en formato PNG, JPG o JPEG"
    )
    
    if firma_file is not None:
        # Procesar firma
        firma_procesada = procesar_firma(firma_file, remover_fondo)
        
        if remover_fondo:
            # Mostrar comparación antes/después
            col1, col2 = st.columns(2)
            with col1:
                st.write("Firma Original")
                st.image(firma_file, width=300)
            with col2:
                st.write("Firma sin fondo")
                st.image(firma_procesada, width=300)
                
            # Opcionalmente mostrar comparador deslizante
            st.write("Comparador deslizante")
            image_comparison(
                img1=Image.open(firma_file),
                img2=Image.open(firma_procesada),
                label1="Original",
                label2="Sin fondo"
            )
        else:
            # Mostrar solo la firma original
            st.image(firma_file, caption="Vista previa de la firma", width=300)
        
        return firma_procesada, True
    
    return None, False

def generar_cotizacion(pdf_file, data):
    # Extraer datos del PDF
    servicio = extraer_nombre_servicio(pdf_file)
    forma_pago = extraer_forma_pago(pdf_file)
    dias = extraer_dias(pdf_file)

    # Actualizar data con los datos extraídos
    data['servicio'] = servicio
    data['armada'] = forma_pago
    data['dias'] = dias

    # Cargar el documento
    template_path = os.path.join(base_dir, 'FormatoCotizacion.docx')
    doc = Document(template_path)

    # Diccionario de reemplazos
    reemplazos = {
        '{{fecha}}': data['fecha'],
        '{{servicio}}': data['servicio'],
        '{{dias}}': data['dias'],
        '{{oferta}}': "{:.2f}".format(data['oferta']),
        '{{armada}}': data['armada'],
        '{{MES}}': data['mes'],
        '{{dni}}': data['dni'],
        '{{nombres}}': data['nombres'],
        '{{ruc}}': data['ruc'],
        '{{telefono}}': data['telefono'],
        '{{correo}}': data['correo'],
        '{{direccion}}': data['direccion'],
        '{{banco}}': data['banco'],
        '{{cuenta}}': data['cuenta'],
        '{{cci}}': data['cci'],
        '{{year}}': str(data['year']),
    }

    def reemplazar_texto(texto, reemplazos):
        for key, value in reemplazos.items():
            texto = texto.replace(key, str(value))
        return texto

    def procesar_parrafo(paragraph):
        if '{{firma}}' in paragraph.text:
            # Manejar la firma como antes
            p = paragraph._element
            p.clear_content()
            run = paragraph.add_run()
            data['firma'].seek(0)
            run.add_picture(BytesIO(data['firma'].read()), height=Cm(1.91))
        else:
            # Concatenar todo el texto de los runs en el párrafo
            full_text = ''
            formatting = []
            for run in paragraph.runs:
                full_text += run.text
                formatting.append({
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    # No guardamos font_name ni font_size
                    'font_color': run.font.color.rgb
                })

            # Reemplazar los marcadores de posición en el texto completo
            new_full_text = reemplazar_texto(full_text, reemplazos)

            # Borrar los runs existentes
            for run in paragraph.runs:
                run.text = ''

            # Crear un nuevo run con el texto reemplazado
            run = paragraph.add_run(new_full_text)
            # Aplicar el formato del primer run original
            if formatting:
                fmt = formatting[0]
                run.bold = fmt['bold']
                run.italic = fmt['italic']
                run.underline = fmt['underline']
                run.font.color.rgb = fmt['font_color']
            else:
                # Valores por defecto si no hay formato original
                run.bold = False
                run.italic = False
                run.underline = False

            # Establecer la fuente a Arial 11
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    # Procesar todos los párrafos en el documento principal
    for paragraph in doc.paragraphs:
        procesar_parrafo(paragraph)

    # Función recursiva para procesar tablas anidadas
    def procesar_tabla(tabla):
        for row in tabla.rows:
            for cell in row.cells:
                # Procesar párrafos dentro de la celda
                for paragraph in cell.paragraphs:
                    procesar_parrafo(paragraph)

                # Procesar tablas anidadas dentro de la celda
                for tabla_anidada in cell.tables:
                    procesar_tabla(tabla_anidada)

    # Procesar todas las tablas en el documento
    for tabla in doc.tables:
        procesar_tabla(tabla)

    # Guardar el documento modificado en un BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def generar_cci(banco, cuenta):
    if not banco or not cuenta or banco == "Otros":
        return ""
    
    cuenta_limpia = cuenta.replace("-", "")
    cci_map = {
        "BCP": "002" + cuenta_limpia + "13",
        "Interbank": "003" + cuenta_limpia + "43",
        "Scotiabank": "00936020" + cuenta_limpia + "95",
        "Banco de la Nación": "0187810" + cuenta_limpia + "55",
        "BanBif": "0386501" + cuenta_limpia + "83"
    }
    return cci_map.get(banco, "")

def main():
    st.set_page_config(
        page_title="Genera tu Cotización",
        page_icon="🎣",
        layout="wide"
    )
    # Inicializar variables de estado para la ubicación
    if 'zoom' not in st.session_state:
        st.session_state['zoom'] = 13
    if 'lat' not in st.session_state:
        st.session_state['lat'] = None
    if 'lon' not in st.session_state:
        st.session_state['lon'] = None
    if 'direccion' not in st.session_state:
        st.session_state['direccion'] = ''

    # Sección de carga de TDR
    st.header("Sube tu TDR (PDF)")
    pdf_file = st.file_uploader("Selecciona tu archivo PDF", type=["pdf"])

    # Sección de firma
    firma_procesada, firma_cargada = mostrar_seccion_firma()

    # Inicializar variables de estado si no existen
    if 'form_data' not in st.session_state:
        st.session_state.form_data = {
            'dni': '',
            'nombres': '',
            'ruc': '',
            'telefono': '',
            'correo': '',
            'direccion': '',
            'banco': '',
            'cuenta': '',
            'cci': '',
            'oferta': 0.0
        }

    # Sección de datos personales
    st.header("Datos Personales")

    # DNI y datos de SUNAT
    dni = st.text_input("Introduce tu DNI", max_chars=8, key='dni_input')
    if dni and len(dni) == 8:
        nombres, ruc = obtener_datos_sunat(dni)
        if nombres:
            st.success(f"Nombres: {nombres}")
            st.success(f"RUC: {ruc}")
            # Guardar en session_state
            st.session_state.form_data.update({
                'dni': dni,
                'nombres': nombres,
                'ruc': ruc
            })

    # Información de contacto
    col1, col2 = st.columns(2)

    with col1:
        telefono = st.text_input("Teléfono", key='telefono_input')
        if telefono:
            st.session_state.form_data['telefono'] = telefono

    with col2:
        correo = st.text_input("Correo electrónico", key='correo_input')
        if correo:
            st.session_state.form_data['correo'] = correo

    st.subheader("Dirección")
    col1, col2 = st.columns([1, 1])

    with col1:
        if st.button("Obtener ubicación actual"):
            get_geolocation('geo_loc')
            st.write("Obteniendo ubicación...")

    # Recuperar la ubicación desde st.session_state después de la llamada
    if 'geo_loc' in st.session_state:
        loc = st.session_state['geo_loc']
        if loc and 'coords' in loc:
            st.session_state.lat = loc['coords']['latitude']
            st.session_state.lon = loc['coords']['longitude']
            direccion = obtener_direccion_desde_coordenadas(
                st.session_state.lat,
                st.session_state.lon
            )
            if direccion:
                st.session_state.direccion = direccion

    # Un solo campo de dirección fuera de las columnas
    direccion_input = st.text_input(
        "Dirección",
        value=st.session_state.get('direccion', ''),
        key="direccion_input"
    )
    # Actualizar el estado con el valor del input
    st.session_state.direccion = direccion_input

    with col2:
        # Mostrar el mapa con la ubicación si está disponible y el zoom actual
        mapa = crear_mapa(
            lat=st.session_state['lat'],
            lon=st.session_state['lon'],
            zoom=st.session_state['zoom']
        )
        mapa_data = st_folium(
            mapa,
            height=300,
            width=None,
            returned_objects=["last_clicked", "zoom"]
        )

        # Actualizar ubicación cuando se hace clic en el mapa
        if mapa_data["last_clicked"]:
            clicked_lat = mapa_data["last_clicked"]["lat"]
            clicked_lng = mapa_data["last_clicked"]["lng"]

            # Guardar el zoom actual antes de actualizar
            if mapa_data.get("zoom"):
                st.session_state['zoom'] = mapa_data["zoom"]

            # Actualizar estado
            st.session_state['lat'] = clicked_lat
            st.session_state['lon'] = clicked_lng
            nueva_direccion = obtener_direccion_desde_coordenadas(clicked_lat, clicked_lng)
            if nueva_direccion:
                st.session_state['direccion'] = nueva_direccion
                st.rerun()
        # Actualizar el zoom incluso si no se hace clic
        elif mapa_data.get("zoom"):
            st.session_state['zoom'] = mapa_data["zoom"]

    # Información bancaria
    st.header("Información Bancaria")
    banco_seleccionado = st.selectbox(
        "Selecciona tu banco",
        ["BCP", "Interbank", "Scotiabank", "Banco de la Nación", "BanBif", "Otros"],
        key='banco_input'
    )
    if banco_seleccionado:
        st.session_state.form_data['banco'] = banco_seleccionado

    cuenta = st.text_input("Ingresa tu cuenta", max_chars=20, key='cuenta_input')
    if cuenta:
        st.session_state.form_data['cuenta'] = cuenta

    # Generar y mostrar CCI
    cci = st.text_input("CCI (editable)", value=generar_cci(banco_seleccionado, cuenta), key='cci_input')
    if cci:
        st.session_state.form_data['cci'] = cci
        
    # Sección de oferta económica
    st.header("Oferta Económica")
    
    # Extraer días del PDF si está disponible
    dias = "30"  # Valor por defecto
    if pdf_file:
        dias = extraer_dias(pdf_file)
    
    # Obtener el valor sugerido basado en los días
    valor_sugerido = obtener_valor_sugerido(dias)
    
    col1, col2 = st.columns([3, 1])
    with col1:
        oferta_total = st.number_input(
            "OFERTA TOTAL (S/)",
            min_value=0.0,
            value=valor_sugerido,  # Valor sugerido dinámico
            step=10.0,
            format="%.2f",
            help=f"Valor sugerido: S/ {valor_sugerido:,.2f} para {dias} días. Puedes ajustar el monto usando las flechas (±10) o ingresando directamente el valor deseado."
        )
    
    with col2:
        st.markdown("""
        <style>
        .small-font {
            font-size: 0.9em;
            color: #666;
        }
        </style>
        """, unsafe_allow_html=True)
        
        st.markdown(
            f'<p class="small-font">Sugerido: S/ {valor_sugerido:,.2f}<br>Incrementos: ± S/ 10.00</p>', 
            unsafe_allow_html=True
        )
    
    # Mostrar el valor ingresado con formato de moneda
    if oferta_total > 0:
        st.write(f"Monto ingresado: S/ {oferta_total:,.2f}")
    
    st.title("Descarga de Constancias")
    # Verificación inicial
    verificar_configuracion_selenium()
    if st.button("Descargar Constancias"):
        with st.spinner("Descargando constancias..."):
            try:
                # Llamada a función de debug
                resultados = descargar_constancias_debug(ruc, dni)
                
                if resultados:
                    st.success("Descarga completada")
                    st.write("Archivos descargados:", resultados)
                else:
                    st.error("No se pudieron descargar las constancias")
            
            except Exception as e:
                st.error(f"Error en la descarga: {e}")
                
    # Botón de envío
    if st.button("Generar cotizacion"):
        if not all([pdf_file, firma_cargada, dni, st.session_state.direccion, telefono, correo, banco_seleccionado, cuenta, cci, oferta_total]):
            st.error("Por favor, complete todos los campos requeridos.")
        else:
            # Obtener datos de SUNAT
            nombres, ruc = obtener_datos_sunat(dni)
            if not nombres:
                st.error("No se pudo obtener datos de SUNAT. Verifica el DNI ingresado.")
            else:
                # Formatear la fecha actual en español
                meses = {
                    "January": "enero", "February": "febrero", "March": "marzo", "April": "abril",
                    "May": "mayo", "June": "junio", "July": "julio", "August": "agosto",
                    "September": "setiembre", "October": "octubre", "November": "noviembre", "December": "diciembre"
                }
                fecha_actual = datetime.now()
                mes_actual = meses[fecha_actual.strftime("%B")].upper()
                fecha_formateada = f"{fecha_actual.day} de {meses[fecha_actual.strftime('%B')]} de {fecha_actual.year}"
                output_directory = os.path.join(base_dir, 'temp_downloads')
                os.makedirs(output_directory, exist_ok=True)

                # Preparar datos para generar la cotización
                data = {
                    'dni': dni,
                    'nombres': nombres,
                    'ruc': ruc,
                    'telefono': telefono,
                    'correo': correo,
                    'direccion': st.session_state.direccion,
                    'banco': banco_seleccionado,
                    'cuenta': cuenta,
                    'cci': cci,
                    'oferta': oferta_total,
                    'fecha': fecha_formateada,
                    'year': fecha_actual.year,
                    'mes': mes_actual,
                    'firma': firma_procesada,
                }

                # Generar la cotización
                doc_io = generar_cotizacion(pdf_file, data)

                # Combinar PDFs de constancias
                combinar_pdfs(output_directory, '5. RNP, RUC, RNSSC.pdf')

                # Crear un archivo ZIP en memoria
                zip_io = BytesIO()
                with zipfile.ZipFile(zip_io, mode='w', compression=zipfile.ZIP_DEFLATED) as zipf:
                    # Agregar el documento de cotización
                    zipf.writestr('Formato de Cotización.docx', doc_io.getvalue())
                    
                    # Agregar la firma
                    firma_procesada.seek(0)  # Reiniciar el puntero del archivo
                    zipf.writestr('Firma.png', firma_procesada.getvalue())
                    
                    # Agregar el TDR original
                    pdf_file.seek(0)  # Reiniciar el puntero del archivo
                    zipf.writestr('6. Copia de Terminos de Referencia.pdf', pdf_file.getvalue())
                    
                    # Agregar el PDF combinado de constancias
                    constancias_path = os.path.join(output_directory, '5. RNP, RUC, RNSSC.pdf')
                    if os.path.exists(constancias_path):
                        with open(constancias_path, 'rb') as constancias_file:
                            zipf.writestr('5. RNP, RUC, RNSSC.pdf', constancias_file.read())

                zip_io.seek(0)
                st.success("¡Cotización generada correctamente!")

                # Botón para descargar el ZIP
                st.download_button(
                    label="Descargar Todos los Archivos Generados (ZIP)",
                    data=zip_io.getvalue(),
                    file_name="cotizacion.zip",
                    mime="application/zip",
                )
    
if __name__ == "__main__":
    main()
